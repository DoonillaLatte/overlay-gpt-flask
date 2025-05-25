import base64
from io import BytesIO
from docx import Document
import logging
from typing import Dict, Any, List, Optional

logger = logging.getLogger(__name__)

class WordService:
    @staticmethod
    def read_docx_data(base64_data: str) -> Optional[Document]:
        """
        base64로 인코딩된 Word 데이터를 읽어 Document 객체로 변환합니다.
        
        Args:
            base64_data (str): base64로 인코딩된 Word 파일 데이터
            
        Returns:
            Document: Word 문서 객체
            None: 에러 발생 시
        """
        try:
            # base64 데이터를 디코딩
            docx_data = base64.b64decode(base64_data)
            
            # BytesIO 객체로 변환
            docx_file = BytesIO(docx_data)
            
            # Word 파일 읽기
            document = Document(docx_file)
            
            return document
        except Exception as e:
            logger.error(f"Word 파일 읽기 오류: {str(e)}")
            return None
    
    @staticmethod
    def get_docx_summary(document: Document) -> Optional[Dict[str, Any]]:
        """
        Word 문서의 기본적인 요약 정보를 반환합니다.
        
        Args:
            document (Document): Word 문서 객체
            
        Returns:
            dict: 데이터 요약 정보
        """
        if document is None:
            return None
            
        try:
            # 단락 스타일 정보 수집
            paragraphs_info = []
            for paragraph in document.paragraphs:
                paragraph_info = {
                    'text': paragraph.text,
                    'style_name': paragraph.style.name,
                    'alignment': paragraph.alignment,
                    'runs': []
                }
                
                # 각 run의 스타일 정보 수집
                for run in paragraph.runs:
                    run_info = {
                        'text': run.text,
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline,
                        'font_name': run.font.name,
                        'font_size': run.font.size,
                        'color': run.font.color.rgb if run.font.color else None
                    }
                    paragraph_info['runs'].append(run_info)
                
                paragraphs_info.append(paragraph_info)
            
            # 표 스타일 정보 수집
            tables_info = []
            for table in document.tables:
                table_info = {
                    'rows': len(table.rows),
                    'columns': len(table.columns),
                    'style': table.style.name
                }
                tables_info.append(table_info)
            
            summary = {
                'paragraphs_count': len(document.paragraphs),
                'tables_count': len(document.tables),
                'paragraphs_info': paragraphs_info,
                'tables_info': tables_info,
                'core_properties': {
                    'author': document.core_properties.author,
                    'title': document.core_properties.title,
                    'subject': document.core_properties.subject,
                    'created': document.core_properties.created,
                    'modified': document.core_properties.modified
                }
            }
            
            return summary
        except Exception as e:
            logger.error(f"Word 데이터 요약 생성 중 오류 발생: {str(e)}")
            return None

    @staticmethod
    def extract_style_definitions(document: Document) -> Optional[Dict[str, Any]]:
        """
        Word 문서의 스타일 정의를 추출합니다.
        
        Args:
            document (Document): Word 문서 객체
            
        Returns:
            dict: 스타일 정의 정보
        """
        if document is None:
            return None
            
        try:
            # 문단 스타일 정의
            paragraph_styles = {}
            for style in document.styles:
                if style.type == 1:  # WD_STYLE_TYPE.PARAGRAPH
                    style_info = {
                        'name': style.name,
                        'base_style': style.base_style.name if style.base_style else None,
                        'font': {
                            'name': style.font.name if style.font else None,
                            'size': style.font.size if style.font else None,
                            'bold': style.font.bold if style.font else None,
                            'italic': style.font.italic if style.font else None
                        },
                        'paragraph_format': {
                            'alignment': style.paragraph_format.alignment if style.paragraph_format else None,
                            'line_spacing': style.paragraph_format.line_spacing if style.paragraph_format else None,
                            'space_before': style.paragraph_format.space_before if style.paragraph_format else None,
                            'space_after': style.paragraph_format.space_after if style.paragraph_format else None
                        }
                    }
                    paragraph_styles[style.name] = style_info
            
            # 문자 스타일 정의
            character_styles = {}
            for style in document.styles:
                if style.type == 2:  # WD_STYLE_TYPE.CHARACTER
                    style_info = {
                        'name': style.name,
                        'base_style': style.base_style.name if style.base_style else None,
                        'font': {
                            'name': style.font.name if style.font else None,
                            'size': style.font.size if style.font else None,
                            'bold': style.font.bold if style.font else None,
                            'italic': style.font.italic if style.font else None
                        }
                    }
                    character_styles[style.name] = style_info
            
            # 표 스타일 정의
            table_styles = {}
            for style in document.styles:
                if style.type == 3:  # WD_STYLE_TYPE.TABLE
                    style_info = {
                        'name': style.name,
                        'base_style': style.base_style.name if style.base_style else None
                    }
                    table_styles[style.name] = style_info
            
            style_definitions = {
                'paragraph_styles': paragraph_styles,
                'character_styles': character_styles,
                'table_styles': table_styles
            }
            
            return style_definitions
        except Exception as e:
            logger.error(f"Word 스타일 정의 추출 중 오류 발생: {str(e)}")
            return None

    @staticmethod
    def extract_content(document: Document) -> Optional[str]:
        """
        Word 문서에서 텍스트 내용만 추출합니다.
        
        Args:
            document (Document): Word 문서 객체
            
        Returns:
            str: 추출된 텍스트 내용
            None: 에러 발생 시
        """
        if document is None:
            return None
            
        try:
            content = []
            
            # 문단 내용 추출
            for paragraph in document.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content.append(text)
            
            # 표 내용 추출
            for table in document.tables:
                table_content = []
                for row in table.rows:
                    row_content = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_content.append(cell_text)
                    if row_content:
                        table_content.append(" | ".join(row_content))
                if table_content:
                    content.append("\n".join(table_content))
            
            # 모든 내용을 하나의 문자열로 합침
            return "\n\n".join(content)
            
        except Exception as e:
            logger.error(f"Word 내용 추출 중 오류 발생: {str(e)}")
            return None 